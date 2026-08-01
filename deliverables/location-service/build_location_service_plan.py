from __future__ import annotations

import re
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


HERE = Path(__file__).resolve().parent
SOURCE = HERE / "REAL_QR_FIND_위치기반서비스_사업계획서_원고.md"
OUTPUT = HERE / "REAL_QR_FIND_위치기반서비스_사업계획서_Word편집본.docx"
ASSET_DIR = HERE / "assets"

FONT_NAME = "맑은 고딕"
FONT_REGULAR = Path(r"C:\Windows\Fonts\malgun.ttf")
FONT_BOLD = Path(r"C:\Windows\Fonts\malgunbd.ttf")

NAVY = "123B5D"
BLUE = "2F79A8"
LIGHT_BLUE = "EAF4FA"
PALE_BLUE = "F4F9FC"
MID_BLUE = "C9E1F0"
GREEN = "317A4A"
LIGHT_GREEN = "EAF6ED"
GOLD = "8A6500"
LIGHT_GOLD = "FFF5D8"
RED = "A43A3A"
LIGHT_RED = "FDEEEE"
GRAY = "5F6B76"
LIGHT_GRAY = "F2F4F6"
BORDER = "AAB4BC"
BLACK = "111827"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_twips: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_twips))
    tc_w.set(qn("w:type"), "dxa")


def set_table_width(table, width_twips: int, indent_twips: int = 110) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_twips))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_twips))
    tbl_ind.set(qn("w:type"), "dxa")


def set_table_grid(table, widths: list[int]) -> None:
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)


def repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_run_font(run, size=10.5, bold=False, color=BLACK, italic=False) -> None:
    run.font.name = FONT_NAME
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT_NAME)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT_NAME)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def paragraph_border_bottom(paragraph, color=NAVY, size=14, space=4) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("페이지 ")
    set_run_font(run, 8.5, color=GRAY)
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char_1, instr_text, fld_char_2])


def set_document_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.8)

    normal = doc.styles["Normal"]
    normal.font.name = FONT_NAME
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_NAME)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_NAME)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.2

    for name, size, color, before, after in (
        ("Heading 1", 17, NAVY, 18, 9),
        ("Heading 2", 14, BLUE, 13, 7),
        ("Heading 3", 11.5, NAVY, 9, 4),
    ):
        style = doc.styles[name]
        style.font.name = FONT_NAME
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT_NAME)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_NAME)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if "Caption KR" not in [s.name for s in doc.styles]:
        style = doc.styles.add_style("Caption KR", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = FONT_NAME
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
        style.font.size = Pt(8.5)
        style.font.color.rgb = RGBColor.from_string(GRAY)
        style.paragraph_format.space_before = Pt(3)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER


def configure_section(section, first=False) -> None:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.8)
    section.different_first_page_header_footer = first

    if not first:
        hp = section.header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = hp.add_run("제자리 | 위치기반서비스 사업계획서")
        set_run_font(r, 8.5, bold=True, color=GRAY)
        paragraph_border_bottom(hp, color=MID_BLUE, size=5, space=3)
        add_page_number(section.footer.paragraphs[0])


def add_cover(doc: Document) -> None:
    section = doc.sections[0]
    configure_section(section, first=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(40)
    p.paragraph_format.space_after = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("위치기반서비스사업 신고 준비본")
    set_run_font(r, 11, bold=True, color=BLUE)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(25)
    p.paragraph_format.space_after = Pt(14)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("위치기반서비스\n사업계획서")
    set_run_font(r, 27, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(18)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("제자리 QR 안심 위치공유 서비스\nREAL_QR_FIND")
    set_run_font(r, 14, bold=True, color=BLUE)

    rule = doc.add_paragraph()
    rule.paragraph_format.space_after = Pt(30)
    paragraph_border_bottom(rule, color=BLUE, size=18, space=6)

    rows = [
        ("상호", "제자리"),
        ("대표자", "이진영, 이진선"),
        ("사업자등록번호", "639-58-00963"),
        ("서비스 URL", "https://zezari.family"),
        ("작성기준일", "2026년 8월 1일"),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_width(table, 7200, 140)
    set_table_grid(table, [1900, 5300])
    for i, (label, value) in enumerate(rows):
        for j, text in enumerate((label, value)):
            cell = table.cell(i, j)
            set_cell_margins(cell, 100, 140, 100, 140)
            set_cell_width(cell, 1900 if j == 0 else 5300)
            if j == 0:
                set_cell_shading(cell, LIGHT_BLUE)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(text)
            set_run_font(r, 10.5, bold=j == 0, color=NAVY if j == 0 else BLACK)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(36)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("문서상태: 제출 준비본 - 보완항목 확인 후 제출")
    set_run_font(r, 9.5, bold=True, color=RED)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("본 문서는 현재 서비스 구현과 운영자료를 기준으로 작성되었습니다.")
    set_run_font(r, 8.5, color=GRAY)


def add_front_matter(doc: Document) -> None:
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(section)

    p = doc.add_paragraph("문서관리 및 제출 전 확인", style="Heading 1")
    paragraph_border_bottom(p, color=BLUE, size=10, space=4)
    add_callout(
        doc,
        "이 문서는 완성 원고이나 신고 제출본으로 확정하기 전에 별도 위치정보 동의, 법정 확인자료 원장, 자동 파기, 좌표 저장 보호, 관리자 감사로그를 구현하고 관련 화면·계약·설정 증빙을 첨부해야 합니다.",
        LIGHT_GOLD,
        GOLD,
    )

    rows = [
        ["문서명", "위치기반서비스 사업계획서"],
        ["상호", "제자리"],
        ["서비스명", "제자리 QR 안심 위치공유 서비스(REAL_QR_FIND)"],
        ["작성기준일", "2026년 8월 1일"],
        ["위치정보관리책임자", "이진선(대표자 지정서 확인 필요)"],
        ["문의", "general@zezari.com / 1668-1290"],
        ["검토기준", "현행 소스코드, 운영 산출물, 2026. 5. 18. 시행 보호조치 기준"],
    ]
    add_table(doc, rows)

    doc.add_heading("목차", level=1)
    toc = [
        "1. 사업자 현황",
        "2. 서비스 내용",
        "3. 위치정보의 보호조치 내역",
        "4. 주요설비 내용 및 설치 장소",
        "별첨 1. 위치정보의 관리지침",
        "별첨 2. 기술적 보호조치 증빙자료 목록",
        "별첨 3. 사업자 확인서류 목록",
        "제출 전 최종 점검표",
        "참고 법령 및 자료",
    ]
    for item in toc:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.4)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(item)
        set_run_font(r, 10.5, bold=item[0].isdigit() or item.startswith("별첨"), color=NAVY)


def add_callout(doc: Document, text: str, fill=LIGHT_BLUE, accent=BLUE) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_width(table, 9640, 180)
    set_table_grid(table, [9640])
    cell = table.cell(0, 0)
    set_cell_width(cell, 9640)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, 120, 180, 120, 180)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "24")
    left.set(qn("w:color"), accent)
    tc_borders.append(left)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    set_run_font(r, 9.5, bold=True, color=NAVY if accent != RED else RED)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def column_widths(count: int) -> list[int]:
    total = 9640
    if count == 2:
        return [2500, 7140]
    if count == 3:
        return [1800, 3820, 4020]
    if count == 4:
        return [1450, 2500, 1750, 3940]
    return [total // count for _ in range(count)]


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    widths = column_widths(cols)
    set_table_width(table, 9640, 110)
    set_table_grid(table, widths)
    for i, row in enumerate(rows):
        prevent_row_split(table.rows[i])
        for j, text in enumerate(row):
            cell = table.cell(i, j)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            set_cell_width(cell, widths[j])
            if i == 0:
                set_cell_shading(cell, NAVY)
            elif cols == 2 and j == 0:
                set_cell_shading(cell, LIGHT_BLUE)
            elif i % 2 == 0:
                set_cell_shading(cell, PALE_BLUE)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.08
            r = p.add_run(text)
            set_run_font(
                r,
                8.5 if cols >= 3 else 9,
                bold=(i == 0 or (cols == 2 and j == 0)),
                color=WHITE if i == 0 else (NAVY if cols == 2 and j == 0 else BLACK),
            )
    repeat_table_header(table.rows[0])
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)


def add_body_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.2
    add_formatted_text(p, text)


def add_formatted_text(paragraph, text: str, size=10.5) -> None:
    parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, size - 0.5, bold=True, color=NAVY)
        elif part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, size, bold=True, color=BLACK)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, size, color=BLACK)


def add_list_item(doc: Document, text: str, ordered=False) -> None:
    p = doc.add_paragraph(style="List Number" if ordered else "List Bullet")
    p.paragraph_format.left_indent = Cm(0.75)
    p.paragraph_format.first_line_indent = Cm(-0.35)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    add_formatted_text(p, text, 10)


def font(size: int, bold=False):
    path = FONT_BOLD if bold and FONT_BOLD.exists() else FONT_REGULAR
    return ImageFont.truetype(str(path), size=size)


def fit_lines(draw, text: str, max_width: int, fnt, max_lines=4) -> list[str]:
    words = text.split()
    lines: list[str] = []
    current = ""
    for word in words:
        trial = f"{current} {word}".strip()
        if draw.textbbox((0, 0), trial, font=fnt)[2] <= max_width:
            current = trial
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    if len(lines) > max_lines:
        lines = lines[:max_lines]
        lines[-1] = lines[-1].rstrip(".") + "..."
    return lines


def draw_box(draw, xy, title, body="", fill="#F4F9FC", outline="#2F79A8", title_color="#123B5D"):
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=18, fill=fill, outline=outline, width=3)
    tf = font(28, True)
    bf = font(21)
    title_lines = fit_lines(draw, title, x2 - x1 - 40, tf, 2)
    ty = y1 + 22
    for line in title_lines:
        box = draw.textbbox((0, 0), line, font=tf)
        draw.text(((x1 + x2 - (box[2] - box[0])) / 2, ty), line, font=tf, fill=title_color)
        ty += 36
    if body:
        ty += 10
        for line in fit_lines(draw, body, x2 - x1 - 46, bf, 5):
            box = draw.textbbox((0, 0), line, font=bf)
            draw.text(((x1 + x2 - (box[2] - box[0])) / 2, ty), line, font=bf, fill="#364152")
            ty += 30


def arrow(draw, start, end, color="#2F79A8"):
    draw.line([start, end], fill=color, width=5)
    x2, y2 = end
    x1, y1 = start
    dx, dy = x2 - x1, y2 - y1
    length = max((dx * dx + dy * dy) ** 0.5, 1)
    ux, uy = dx / length, dy / length
    px, py = -uy, ux
    p1 = (x2 - ux * 22 + px * 10, y2 - uy * 22 + py * 10)
    p2 = (x2 - ux * 22 - px * 10, y2 - uy * 22 - py * 10)
    draw.polygon([end, p1, p2], fill=color)


def canvas(title: str):
    img = Image.new("RGB", (1500, 820), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    draw.rounded_rectangle((10, 10, 1490, 810), radius=20, fill="#FFFFFF", outline="#AAB4BC", width=3)
    draw.text((55, 32), title, font=font(34, True), fill="#123B5D")
    draw.line((55, 84, 1445, 84), fill="#C9E1F0", width=5)
    return img, draw


def make_diagrams() -> dict[str, Path]:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    outputs: dict[str, Path] = {}

    img, draw = canvas("위치정보 보호 조직")
    draw_box(draw, (540, 125, 960, 250), "대표자", "정책 승인 · 예산 · 최종 의사결정", "#EAF4FA")
    draw_box(draw, (540, 335, 960, 480), "위치정보관리책임자", "보호업무 총괄 · 교육 · 점검 · 사고대응", "#EAF6ED", "#317A4A")
    for x, title, body in (
        (95, "보호담당자", "권한 · 로그 · 증빙 · 내부검사"),
        (540, "위치정보취급자", "승인된 범위의 조회 · 민원 처리"),
        (985, "개발·운영 담당", "배포 · 보안설정 · 복구 · 개선"),
    ):
        draw_box(draw, (x, 600, x + 420, 750), title, body)
    arrow(draw, (750, 250), (750, 335), "#317A4A")
    for x in (305, 750, 1195):
        arrow(draw, (750, 480), (x, 600))
    path = ASSET_DIR / "01_위치정보보호조직.png"
    img.save(path)
    outputs["organization"] = path

    img, draw = canvas("QR 발견 위치공유 서비스 시나리오")
    boxes = [
        (55, "보호자 등록", "관리대상 등록\nQR 활성화"),
        (345, "QR 스캔", "발견자가 공개\n페이지 접속"),
        (635, "별도 동의", "안내 확인 후\n위치권한 허용"),
        (925, "위치 저장", "좌표 검증\n이력 기록"),
        (1215, "보호자 알림", "카카오맵 링크\nWeb Push"),
    ]
    for x, title, body in boxes:
        draw_box(draw, (x, 245, x + 235, 510), title, body, "#F4F9FC")
    for x in (290, 580, 870, 1160):
        arrow(draw, (x, 378), (x + 45, 378))
    draw.text((55, 610), "핵심 원칙", font=font(28, True), fill="#123B5D")
    draw_box(draw, (240, 585, 1425, 735), "지속 추적 없음", "발견자가 버튼을 누른 시점의 위치만 1회 수집하며, 거부 시 위치공유만 제한합니다.", "#FFF5D8", "#8A6500", "#8A6500")
    path = ASSET_DIR / "02_서비스시나리오.png"
    img.save(path)
    outputs["scenario"] = path

    img, draw = canvas("위치정보 데이터 흐름")
    boxes = [
        ((50, 185, 310, 360), "발견자 단말", "브라우저·OS\n위도·경도·정확도"),
        ((390, 185, 650, 360), "Vercel API", "동의·QR·권한\n좌표 범위 검증"),
        ((730, 185, 990, 360), "Turso DB", "위치공유 이력\n이용·제공사실"),
        ((1070, 185, 1330, 360), "Web Push", "암호화 알림\n지도 링크"),
        ((1070, 535, 1330, 710), "보호자", "알림 클릭\n카카오맵 확인"),
        ((730, 535, 990, 710), "권한 관리자", "이력 조회\n감사로그"),
    ]
    for xy, title, body in boxes:
        draw_box(draw, xy, title, body)
    arrow(draw, (310, 272), (390, 272))
    arrow(draw, (650, 272), (730, 272))
    arrow(draw, (990, 272), (1070, 272))
    arrow(draw, (1200, 360), (1200, 535))
    arrow(draw, (860, 360), (860, 535), "#317A4A")
    draw.text((55, 445), "전송구간 TLS", font=font(24, True), fill="#317A4A")
    draw.text((305, 445), "최소수집 · 최소권한 · 보존기간 · 자동파기", font=font(24, True), fill="#123B5D")
    path = ASSET_DIR / "03_데이터흐름도.png"
    img.save(path)
    outputs["dataflow"] = path

    img, draw = canvas("관리적·기술적 보호조치 구조")
    draw_box(draw, (70, 150, 470, 690), "관리적 보호조치", "책임자 지정\n최소권한·권한대장\n취급대장\n연간 교육·자체검사\n위탁관리\n침해사고 대응", "#EAF6ED", "#317A4A", "#317A4A")
    draw_box(draw, (550, 150, 950, 690), "위치정보 처리 원칙", "별도 동의\n목적 제한\n단발성 수집\n확인자료 분리보관\n원본 좌표 자동파기\n이용자 권리보장", "#FFF5D8", "#8A6500", "#8A6500")
    draw_box(draw, (1030, 150, 1430, 690), "기술적 보호조치", "식별·인증·MFA\nHTTPS/TLS\nDB·필드 암호화\n방화벽·입력검증\n접근로그·변경방지\n백신·취약점 점검", "#EAF4FA", "#2F79A8", "#123B5D")
    arrow(draw, (470, 420), (550, 420), "#317A4A")
    arrow(draw, (950, 420), (1030, 420), "#2F79A8")
    path = ASSET_DIR / "04_보호조치구조.png"
    img.save(path)
    outputs["security"] = path

    img, draw = canvas("주요설비 및 설치장소 구성")
    draw_box(draw, (60, 170, 340, 340), "사용자·발견자", "PWA · 모바일 브라우저")
    draw_box(draw, (430, 145, 760, 365), "Vercel", "CDN · HTTPS\nNext.js · Serverless API", "#EAF4FA")
    draw_box(draw, (850, 145, 1180, 365), "Turso/libSQL", "서비스·위치 이력\n백업 · 접근 토큰", "#EAF6ED", "#317A4A")
    draw_box(draw, (1185, 510, 1440, 700), "지도·푸시", "Kakao/Naver Map\nWeb Push")
    draw_box(draw, (430, 510, 760, 700), "GitHub·Vercel 배포", "소스·변경이력\n빌드·롤백")
    draw_box(draw, (60, 510, 340, 700), "관리자 PC", "MFA · 백신\n화면잠금 · 로그점검", "#FFF5D8", "#8A6500", "#8A6500")
    arrow(draw, (340, 255), (430, 255))
    arrow(draw, (760, 255), (850, 255))
    arrow(draw, (1180, 300), (1310, 510))
    arrow(draw, (600, 510), (600, 365))
    arrow(draw, (340, 605), (430, 605), "#8A6500")
    draw.text((850, 555), "클라우드 실제 리전과 계약증빙은\n제출 시점의 대시보드로 확인", font=font(23, True), fill="#A43A3A")
    path = ASSET_DIR / "05_주요설비구성도.png"
    img.save(path)
    outputs["infrastructure"] = path

    return outputs


FIGURE_CAPTIONS = {
    "organization": "그림 1. 위치정보 보호 조직(제출 전 실제 담당자 지정 필요)",
    "scenario": "그림 2. 제자리 QR 안심 위치공유 서비스 시나리오",
    "dataflow": "그림 3. 위치정보의 수집·이용·제공 데이터 흐름",
    "security": "그림 4. 관리적·기술적 보호조치 구조",
    "infrastructure": "그림 5. 위치정보시스템 주요설비 및 설치장소 구성",
}


def add_figure(doc: Document, key: str, figures: dict[str, Path]) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    run.add_picture(str(figures[key]), width=Inches(6.55))
    cp = doc.add_paragraph(FIGURE_CAPTIONS[key], style="Caption KR")
    cp.paragraph_format.keep_with_next = False


def parse_markdown(doc: Document, figures: dict[str, Path]) -> None:
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    start = next(i for i, line in enumerate(lines) if line.strip() == "# 1. 사업자 현황")
    lines = lines[start:]
    i = 0
    first_h1 = True
    paragraph_buffer: list[str] = []

    def flush_paragraph():
        if paragraph_buffer:
            text = " ".join(s.strip() for s in paragraph_buffer).strip()
            if text:
                add_body_paragraph(doc, text)
            paragraph_buffer.clear()

    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()

        if not stripped:
            flush_paragraph()
            i += 1
            continue

        if stripped.startswith("[[FIGURE:"):
            flush_paragraph()
            key = stripped[len("[[FIGURE:") : -2]
            add_figure(doc, key, figures)
            i += 1
            continue

        if stripped.startswith("# "):
            flush_paragraph()
            if not first_h1:
                doc.add_page_break()
            first_h1 = False
            p = doc.add_heading(stripped[2:], level=1)
            paragraph_border_bottom(p, color=BLUE, size=10, space=4)
            i += 1
            continue

        if stripped.startswith("## "):
            flush_paragraph()
            doc.add_heading(stripped[3:], level=2)
            i += 1
            continue

        if stripped.startswith("### "):
            flush_paragraph()
            doc.add_heading(stripped[4:], level=3)
            i += 1
            continue

        if stripped.startswith("> "):
            flush_paragraph()
            text = stripped[2:]
            add_callout(doc, text, LIGHT_GOLD, GOLD)
            i += 1
            continue

        if stripped.startswith("| "):
            flush_paragraph()
            raw_rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                raw = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                raw_rows.append(raw)
                i += 1
            rows = [r for idx, r in enumerate(raw_rows) if idx != 1 or not all(re.fullmatch(r":?-{3,}:?", c) for c in r)]
            add_table(doc, rows)
            continue

        num_match = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if num_match:
            flush_paragraph()
            add_list_item(doc, num_match.group(2), ordered=True)
            i += 1
            continue

        if stripped.startswith("- "):
            flush_paragraph()
            add_list_item(doc, stripped[2:], ordered=False)
            i += 1
            continue

        paragraph_buffer.append(stripped)
        i += 1

    flush_paragraph()


def add_back_matter(doc: Document) -> None:
    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(100)
    r = p.add_run("제출 전 확인")
    set_run_font(r, 22, bold=True, color=NAVY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run("실제 신고 시에는 최신 법령·신고서식과 운영환경을 다시 대조하고,\n미완료 보호조치를 구현한 뒤 증빙을 첨부해 주세요.")
    set_run_font(r, 12, bold=True, color=RED)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("제자리 · general@zezari.com · 1668-1290")
    set_run_font(r, 10, color=GRAY)


def build() -> Path:
    figures = make_diagrams()
    doc = Document()
    set_document_styles(doc)
    add_cover(doc)
    add_front_matter(doc)
    parse_markdown(doc, figures)
    add_back_matter(doc)

    core = doc.core_properties
    core.title = "REAL_QR_FIND 위치기반서비스 사업계획서"
    core.subject = "제자리 QR 안심 위치공유 서비스 위치기반서비스사업 신고 준비본"
    core.author = "제자리"
    core.keywords = "위치기반서비스, 위치정보, REAL_QR_FIND, 제자리, 사업계획서"
    core.comments = "현재 구현과 신고 전 보완항목을 구분한 제출 준비본"

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build())
